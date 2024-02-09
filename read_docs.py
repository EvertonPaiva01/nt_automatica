from docx import Document
from docx.shared import Pt

# Abre o documento .docx
doc = Document("NT XX.2024 - Vistoria Pólo de Carnaval 'Nome'.docx")

# Definir o estilo de parágrafo padrão
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

# Define as informações padrões da nota
nome_polo = input("Digite o nome do Polo: Ex: Marco Zero\n")
logradouro = input("Digite o Logradouro: Ex: Av. Alfredo Lisboa\n")
bairro = input("Digite o Bairro: Ex: Recife\n")
dia = input("Digite a dia da ocorrencia: Ex: 08\n")
hora = input("Digite a hora: Ex: 10:00\n")
nome_tecnico = input("Digite o nome do Técnico responsável: Ex: Thiago Silva\n")
cargo = input("Digite o cargo do técnico responsável: Ex: Gerente Geral de Iluminação Pública\n")
informacoes = input("Digite o número dos itens que apresentaram irregularidades: Ex: 1,2,3\n")
lista_numeros = [int(num) for num in informacoes.split(',')]

visto_info = {
    1:"A infraestrutura não possui ART",

    2:"O projeto elétrico com os equipamentos a serem instalados não foram encaminhado para a EMLURB juntamente a ART do responsável técnico pelas instalações",

    3:"O responsável técnico não estava presente durante a montagem",

    4:"O fornecimento da energia não foi realizado pela Neoenergia PE",
    
    5:"Existem cargas conectadas diretamente a rede elétrica de IP",

    6:"A instalação não possui Disjuntor Termomagnético",

    7:"A instalação não possui DPS",

    8:"A instalação não possui DR",

    9:"O DR instalado não é de alta sensibilidade (30mA)",

    10:"As partes metálicas que são componentes dos sistemas elétricos não estão aterradas",

    11:"As partes metálicas que são próximas as instalações elétricas não estão aterradas",

    12:'O aterramento não está executado com vara Copperweld 5/8"x2,40m, fincada no solo com conector do tipo GAR/GTDU ou solda exotérmica, utilizando cabo de cobre 0,6/1kV de cor verde com bitola mínima de 16mm²',

    13:"O aterramento das massas não foi respeitado",

    14:"Existem cabeamentos elétricos expostos ao tempo, intempéries, passando pelo chão ou amarrados em estruturas metálicas",

    15:"Não foi utilizado um trilho passa cabo protegendo os cabeamentos",

    16:"Existem projetores/refletores ou qualquer tipo de iluminação provisória, em postes metálicos(exclusivos para iluminação pública) ou árvores",

    17:"Os equipamentos usados para a realização das instalações elétricas provisórias não são de boa qualidade",

    18:"Existem elementos (bandeira, barracas, refletores, fiações, entre outros) nos postes e demais equipamentos componentes do sistema de iluminação pública",

    19:"As distâncias para a AT e MT não estão respeitadas(conforme ABNT e Neoenergia)",

    20:"As caixas de passagem subterrânea estão violadas",

    21:"Os medidores da Neoenergia estão violados",

    22:"Existem cordões luminosos ou gambiarras cruzando vias, grades ou elementos metálicos pertencentes às estruturas dos prédios, comércios, residências, etc"
}

# Busca na Tabela
def table_search(search,replace):
    for table in doc.tables:
        # Itera sobre todas as linhas da tabela
        for row in table.rows:
            # Itera sobre todas as células da linha
            for cell in row.cells:
                # Itera sobre todos os parágrafos dentro da célula
                for paragraph in cell.paragraphs:
                    # Verifica se o trecho desejado está no texto do parágrafo
                    if search in paragraph.text:
                        # Substitui "Nome do técnico." pelo nome do trabalhador
                        print(paragraph.text)
                        paragraph.text = paragraph.text.replace(search, replace)
         
def table_drop(search,replace):
    for table in doc.tables:
        # Itera sobre todas as linhas da tabela
        for row in table.rows:
            # Itera sobre todas as células da linha
            for cell in row.cells:
                # Itera sobre todos os parágrafos dentro da célula
                for paragraph in cell.paragraphs:
                    # Verifica se o trecho desejado está no texto do parágrafo
                    if search in paragraph.text:
                        # Substitui "Nome do técnico." pelo nome do trabalhador
                        print(paragraph.text)
                        paragraph.text = paragraph.text.replace(search, replace)

table_search("Nome do Pólo",nome_polo)
#table_search("Nome do Pólo",nome_polo)
table_search("‘Logradouro’ bairro",logradouro + ", ")
table_search("Logradouro",logradouro )
table_search("‘Nome do bairro’",bairro)
table_search("Bairro",bairro)
table_search("XX de fevereiro",dia + " de fevereiro")
table_search("XX/02/2024",dia +"/02/2024")
table_search("XX:XX",hora)
table_search("Nome do técnico",nome_tecnico )
table_search("Nome do Responsável pela vistoria",nome_tecnico )
table_search("Cargo",cargo)

for i in range(1,23):
    if i in lista_numeros and i in visto_info:
        if i < 10:
            indice_vistoria = "0"+str(i)+" - Informações da vistoria."
            print(indice_vistoria)
            table_search(indice_vistoria , visto_info[i]+".")
        else:
            indice_vistoria = str(i)+" - Informações da vistoria."
            print(indice_vistoria)
            table_search(indice_vistoria , visto_info[i]+".")
    else:
        if i < 10:
            indice_vistoria = "0"+str(i)+" - Informações da vistoria."
            print(indice_vistoria)
            table_search(indice_vistoria ,"-------Remover-----")
        else:
            indice_vistoria = str(i)+" - Informações da vistoria."
            table_search(indice_vistoria ,"-------Remover-----")
            
# Salva as alterações no documento
doc.save("NT XX.2024 - Vistoria Pólo de Carnaval "+ nome_polo +".docx")