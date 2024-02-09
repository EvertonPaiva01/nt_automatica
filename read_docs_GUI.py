from docx import Document
import PySimpleGUI as sg
from docx import Document

sg.theme('GrayGrayGray')
layout = [
    [sg.Image(r'C:\nt_automatica\EMLURB.png')],
    [sg.Text('Digite o nome do Polo: Ex: Marco Zero'), sg.InputText(key = 'nome_polo')],
    [sg.Text('Digite o Logradouro: Ex: Av. Alfredo Lisboa\n'), sg.InputText(key = 'logradouro')],
    [sg.Text('Digite o Bairro: Ex: Casa Amarela'), sg.InputText(key = 'bairro')],
    [sg.Text('Digite a dia da ocorrencia: Ex: 08'), sg.InputText(key = 'dia')],
    [sg.Text('Digite a hora: Ex: 10:00'), sg.InputText(key = 'hora')],
    [sg.Text('Digite o nome do Técnico responsável: Ex: Thiago Silva'), sg.InputText(key = 'nome_tecnico')],
    [sg.Text('Digite o cargo do técnico responsável: Ex: Gerente Geral de Iluminação Pública'), sg.InputText(key = 'cargo')],
    [sg.Text('Digite o número dos itens que apresentaram irregularidades: Ex: 1,2,3\n'), sg.InputText(key = 'informações')],
    [sg.Button("Enviar Informações")]
]

doc = Document("NT XX.2024 - Vistoria Pólo de Carnaval 'Nome'.docx")
window  = sg.Window("Gerador automático de Notas Técnicas",layout)

def table_search(search,replace):
    for table in doc.tables:
        # Itera sobre todas as linhas da tabela
        for row in table.rows:
            # Itera sobre todas as células da linha
            for cell in row.cells:
                # Itera sobre todos os parágrafos dentro da célula
                for paragraph in cell.paragraphs:
                    # Verifica se o trecho desejado está no texto do parágrafo
                    if search in paragraph.text:
                        # Substitui "Nome do técnico." pelo nome do trabalhador
                        print(paragraph.text)
                        paragraph.text = paragraph.text.replace(search, replace)

while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED:
        break 

    visto_info = {
        1:"A infraestrutura não possui ART",

        2:"O projeto elétrico com os equipamentos a serem instalados não foram encaminhado para a EMLURB juntamente a ART do responsável técnico pelas instalações",

        3:"O responsável técnico não estava presente durante a montagem",

        4:"O fornecimento da energia não foi realizado pela Neoenergia PE",

        5:"Existem cargas conectadas diretamente a rede elétrica de IP",

        6:"A instalação não Disjuntor Termomagnético",

        7:"A instalação não possui DPS",

        8:"A instalação não possui DR",

        9:"O DR instalado não é de alta sensibilidade (30mA)",

        10:"As partes metálicas que são componentes dos sistemas elétricos não estão aterradas",

        11:"As partes metálicas que são próximas as instalações elétricas não estão aterradas",

        12:"O aterramento não está executado com vara Copperweld 5/8x2,40m, fincada no solo com conector do tipo GAR/GTDU ou solda exotérmica, utilizando cabo de cobre 0,6/1kV de cor verde com bitola mínima de 16mm²",

        13:"O aterramento das massas não foi respeitado",

        14:"Existem cabeamentos elétricos expostos ao tempo, intempéries, passando pelo chão ou amarrados em estruturas metálicas",

        15:"Não foi utilizado um trilho passa cabo protegendo os cabeamentos",

        16:"Existem projetores/refletores ou qualquer tipo de iluminação provisória, em postes metálicos(exclusivos para iluminação pública) ou árvores",

        17:"Os equipamentos usados para a realização das instalações elétricas provisórias não são de boa qualidade",

        18:"Existem elementos (bandeira, barracas, refletores, fiações, entre outros) nos postes e demais equipamentos componentes do sistema de iluminação pública",

        19:"As distâncias para a AT e MT não estão respeitadas(conforme ABNT e Neoenergia)",

        20:"As caixas de passagem subterrânea estão violadas",

        21:"Os medidores da Neoenergia estão violados",

        22:"Existem cordões luminosos ou gambiarras cruzando vias, grades ou elementos metálicos pertencentes às estruturas dos prédios, comércios, residências, etc"
    }

    if event == 'Enviar Informações':
        nome_polo = values['nome_polo']
        logradouro = values['logradouro']
        bairro = values['bairro']
        dia = values['dia']
        hora = values['hora']
        nome_tecnico = values['nome_tecnico']
        cargo = values['cargo']
        informações = values['informações']
        lista_numeros = [int(num) for num in informações.split(',')]
        
        print(nome_polo, logradouro, bairro, dia, hora, nome_tecnico, cargo)

        table_search("Nome do Pólo", values['nome_polo'])
        table_search("Logradouro",logradouro )
        table_search("Bairro",bairro)
        table_search("XX de fevereiro",dia)
        table_search("XX/02/2024",dia +"/02/2024")
        table_search("XX:XX",hora)
        table_search("Nome do técnico",nome_tecnico )
        table_search("Nome do Responsável pela vistoria",nome_tecnico )
        table_search("Cargo",cargo)

        for i in visto_info:
            if i in lista_numeros:
                if i < 10:
                    indice_vistoria = "0"+str(i)+" - Informações da vistoria."
                    print(indice_vistoria)
                    table_search(indice_vistoria , visto_info[i]+".")
                else:
                    indice_vistoria = str(i)+" - Informações da vistoria."
                    print(indice_vistoria)
                    table_search(indice_vistoria , visto_info[i]+".")
            else:
                if i < 10:
                    indice_vistoria = "0"+str(i)+" - Informações da vistoria."
                    print(indice_vistoria)
                    table_search(indice_vistoria ,"-------Remover-----")
                else:    
                    indice_vistoria = str(i)+" - Informações da vistoria."
                    print(indice_vistoria)
                    table_search(indice_vistoria ,"-------Remover-----")

        window.close()

    doc.save("seu_documento_modificado.docx")